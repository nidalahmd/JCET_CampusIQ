import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from pypdf import PdfReader


@dataclass
class ParsedSection:
    content: str
    page_number: int | None = None
    section_title: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    sections: list[ParsedSection]
    total_pages: int = 1
    file_type: str = "txt"
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(s.content for s in self.sections if s.content.strip())


def parse_pdf(file_path: Path) -> ParsedDocument:
    reader = PdfReader(str(file_path))
    total_pages = len(reader.pages)
    sections: list[ParsedSection] = []

    for page_idx, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if not page_text.strip():
            continue

        # Split by potential headings (lines that are short and title-case/uppercase)
        lines = page_text.splitlines()
        current_heading: str | None = f"Page {page_idx}"
        current_block: list[str] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Heading heuristic: line length < 70, no trailing period, title-case or uppercase or starts with number/section
            is_heading = (
                len(stripped) < 70
                and not stripped.endswith(".")
                and (
                    stripped.isupper()
                    or stripped.istitle()
                    or bool(re.match(r"^(\d+(\.\d+)*|[A-Z]\.)\s+", stripped))
                )
            )

            if is_heading and len(current_block) > 0:
                block_content = "\n".join(current_block).strip()
                if block_content:
                    sections.append(
                        ParsedSection(
                            content=block_content,
                            page_number=page_idx,
                            section_title=current_heading,
                        )
                    )
                current_heading = stripped
                current_block = []
            else:
                current_block.append(stripped)

        if current_block:
            block_content = "\n".join(current_block).strip()
            if block_content:
                sections.append(
                    ParsedSection(
                        content=block_content,
                        page_number=page_idx,
                        section_title=current_heading,
                    )
                )

    return ParsedDocument(
        sections=sections,
        total_pages=total_pages,
        file_type="pdf",
        metadata={"page_count": total_pages},
    )


def parse_docx(file_path: Path) -> ParsedDocument:
    import docx

    doc = docx.Document(str(file_path))
    sections: list[ParsedSection] = []
    current_heading: str | None = "Document Header"
    current_block: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else ""
        is_heading = "Heading" in style_name or "Title" in style_name

        if is_heading:
            if current_block:
                block_content = "\n".join(current_block).strip()
                if block_content:
                    sections.append(
                        ParsedSection(
                            content=block_content,
                            page_number=1,
                            section_title=current_heading,
                        )
                    )
                current_block = []
            current_heading = text
        else:
            current_block.append(text)

    # Process tables if any
    for table in doc.tables:
        table_rows: list[str] = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_rows.append(row_text)
        if table_rows:
            current_block.append("\n".join(table_rows))

    if current_block:
        block_content = "\n".join(current_block).strip()
        if block_content:
            sections.append(
                ParsedSection(
                    content=block_content,
                    page_number=1,
                    section_title=current_heading,
                )
            )

    return ParsedDocument(
        sections=sections,
        total_pages=1,
        file_type="docx",
    )


def parse_markdown(file_path: Path) -> ParsedDocument:
    raw_content = file_path.read_text(encoding="utf-8", errors="replace")
    sections: list[ParsedSection] = []
    lines = raw_content.splitlines()

    current_heading: str | None = "Introduction"
    current_block: list[str] = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#"):
            # Markdown header
            header_match = re.match(r"^#+\s+(.+)$", stripped)
            if header_match:
                if current_block:
                    block_content = "\n".join(current_block).strip()
                    if block_content:
                        sections.append(
                            ParsedSection(
                                content=block_content,
                                page_number=1,
                                section_title=current_heading,
                            )
                        )
                    current_block = []
                current_heading = header_match.group(1)
                continue

        current_block.append(line)

    if current_block:
        block_content = "\n".join(current_block).strip()
        if block_content:
            sections.append(
                ParsedSection(
                    content=block_content,
                    page_number=1,
                    section_title=current_heading,
                )
            )

    return ParsedDocument(
        sections=sections,
        total_pages=1,
        file_type="md",
    )


def parse_text(file_path: Path) -> ParsedDocument:
    raw_content = file_path.read_text(encoding="utf-8", errors="replace")
    sections: list[ParsedSection] = []
    paragraphs = raw_content.split("\n\n")

    current_heading: str | None = "Document Content"
    for para in paragraphs:
        cleaned = para.strip()
        if not cleaned:
            continue

        # Check if single line is a header
        lines = cleaned.splitlines()
        if len(lines) > 1 and len(lines[0].strip()) < 60 and lines[0].strip().isupper():
            current_heading = lines[0].strip()
            cleaned = "\n".join(lines[1:]).strip()

        if cleaned:
            sections.append(
                ParsedSection(
                    content=cleaned,
                    page_number=1,
                    section_title=current_heading,
                )
            )

    if not sections and raw_content.strip():
        sections.append(
            ParsedSection(
                content=raw_content.strip(),
                page_number=1,
                section_title="Document Content",
            )
        )

    return ParsedDocument(
        sections=sections,
        total_pages=1,
        file_type="txt",
    )


def parse_document(file_path: Path, file_type: str | None = None) -> ParsedDocument:
    ext = (file_type or file_path.suffix.lower().lstrip(".")).lower()
    if ext == "pdf":
        return parse_pdf(file_path)
    elif ext in ("docx", "doc"):
        return parse_docx(file_path)
    elif ext in ("md", "markdown"):
        return parse_markdown(file_path)
    else:
        return parse_text(file_path)
